from flask import Flask, request, send_file
import docx
from docx import Document
from docx.oxml.ns import qn
import datetime
import json
doc_file_name = ""
DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def makepara(text, style=3):
    para = doc.add_paragraph()
    pw = para.add_run(text)
    pw.font.name = '游ゴシック Medium'
    pw._element.rPr.rFonts.set(qn('w:eastAsia'), pw.font.name)
    if style == 1:
        pw.font.size = 25*12700
    elif style == 2:
        pw.font.size = 20*12700
    else:
        pw.font.size = 15*12700


def main(request_data, req_id):
    global doc
    doc = docx.Document()
    dt = datetime.datetime.now()
    print('main funciton was run')
    #wordデータの生成・書き込み
    makepara('実験計画書', style=1)
    makepara(f'表題　{request_data["hyoudai"]}', style=2)
    makepara(f'実験予定日　{request_data["yoteibi"]}')
    makepara(f'実験実施日　{request_data["jissibi"]}')
    makepara(f'記入者　{request_data["kinyuusha"]}')
    makepara(f'参加者　{request_data["sankasha"]}')
    makepara(f'目的　{request_data["mokuteki"]}')
    makepara(f'器具　{request_data["kigu"]}')
    makepara(f'薬品　{request_data["yakuhin"]}')
    makepara(f'手順・予想・注意　{request_data["tejun"]}')
    makepara(f'結果　{request_data["kekka"]}')
    makepara(f'考察・課題　{request_data["kousatu"]}')
    f=open('filedata.json', "r")
    print(f)
    #json = f.read()
    #print(json)
    filedata_content = json.loads(f.read().encode('unicode-escape').decode('unicode-escape'))
    json_dict = filedata_content["filedata"]
    print(type(json_dict), json_dict)
    print(list(json_dict.keys()))

    #ログデータの生成・filedata.jsonへ書き込み
    if request_data["hyoudai"] in list(json_dict.keys()):
        print('true')
        last_file_num = len(json_dict[request_data["hyoudai"]])
        print('true2')
        request_data_hyoudai=request_data["hyoudai"]
        print('true3')
        print('true4')
        #request_data_hyoudai["1"] = {}
        #request_data_hyoudai_1 = request_data_hyoudai["1"]
        print(request_data["kinyuusha"])
        print("last_file_num:", last_file_num)
        json_dict[str(request_data_hyoudai)][str(last_file_num+1)]={"user": str(request_data["kinyuusha"]), "date": f"{dt.year}/{dt.month}/{dt.day}/{dt.hour}/{dt.minute}/{dt.second}", "content": request_data}
        print("true5")
        #json_dict[request_data_hyoudai][str(last_file_num+1)]["date"]=
        print(request_data_hyoudai)
        print(json_dict)
        doc_file_name = request_data["hyoudai"]+str(last_file_num+1)
        doc.save(f'docx_files/{request_data["hyoudai"]}{last_file_num+1}.docx')
    else:
        print('true')
        print('true2')
        request_data_hyoudai=request_data["hyoudai"]
        print('true3')
        print('true4')
        #request_data_hyoudai["1"] = {}
        #request_data_hyoudai_1 = request_data_hyoudai["1"]
        print(request_data["kinyuusha"])
        #print("last_file_num:", last_file_num)
        json_dict[request_data_hyoudai]={}
        json_dict[str(request_data_hyoudai)][str(1)]={"user": str(request_data["kinyuusha"]), "date": f"{dt.year}/{dt.month}/{dt.day}/{dt.hour}/{dt.minute}/{dt.second}", "content": request_data}
        print("true5")
        #json_dict[request_data_hyoudai][str(last_file_num+1)]["date"]=
        print(request_data_hyoudai)
        print(json_dict)
        doc_file_name = request_data["hyoudai"]+str(1)
        doc.save(f'docx_files/{request_data["hyoudai"]}{"1"}.docx')
    f.close()
    f_write = open('filedata.json', "w")
    filedata_content["filedata"]=json_dict
    filedata_content["id_name"][str(req_id)]=doc_file_name
    print("f_write")
    json.dump(filedata_content, f_write)
    f_write.close()
    return doc_file_name

def data_import():
    with open('filedata.json') as f:
        json_dict = json.loads(f.read().encode('unicode-escape').decode('unicode-escape'))["filedata"]
    return json.dumps(json_dict)


app = Flask(__name__, static_folder="static")

@app.route('/getid', methods=["GET"])
def getid():
    with open("filedata.json", "r") as f:
        file_content=json.loads(f.read().encode('unicode-escape').decode('unicode-escape'))
        lastid=int(file_content["lastid"])
        file_content["lastid"]=str(lastid+1)
    with open('filedata.json', "w") as f2:
        f2.write(json.dumps(file_content))
    return json.dumps({"id": str(lastid)})

@app.route('/sendrequest', methods=["POST"])
def get_request():
    print('get_request function was run')
    print('dict:')
    request_data = json.loads(str(request.get_data().decode('utf-8')))
    print(request_data)
    print("2")
    if request_data["mode"]=="main":
        filename = main(request_data["content"], request_data["id"])
        DownloadFileName = filename
        downloadFile = f"docx_files/{filename}.docx"
        print(f"filename:{downloadFile}")
        return send_file(downloadFile, mimetype=DOCX_MIMETYPE)
    elif request_data["mode"]=="import":
        return data_import()
    elif request_data["mode"]=="getfilename":
        with open("filedata.json", "r") as f:
            file_content=json.loads(f.read().encode('unicode-escape').decode('unicode-escape'))["id_name"]
            docx_file_name = file_content[str(request_data["id"])]
        print(docx_file_name)
        return json.dumps({"filename": str(docx_file_name)})


@app.route('/')
def post_html():
    with open('index.html', "r") as html:
        html_data=html.read()
        print('html: ', html_data)
    print('response:', html_data)
    return html_data

@app.route('/html/style.css')
def style_css():
    with open('style.css', "r") as css:
        css_data = css.read()
    return css_data


@app.after_request
def after_request(response):
    response.headers['Access-Control-Allow-Origin'] = "*" 
    return response

if __name__ == "__main__":
    app.run()